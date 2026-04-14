from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "strategy" / "servicemen-app"
OUTPUT_DIR = SOURCE_DIR / "docx"

SOURCE_FILES = [
    "persona-clarification-document.md",
    "mvp-scope-prioritisation-document.md",
    "role-by-role-feature-requirements-matrix.md",
    "whatsapp-journey-specification.md",
]


def clean_cell(text: str) -> str:
    return text.strip().strip(":").strip()


def is_table_separator(line: str) -> bool:
    if "|" not in line:
        return False
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    if not cells:
        return False
    return all(cell and set(cell) <= {"-", ":"} for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].rstrip("\n"))
        i += 1

    if len(table_lines) < 2 or not is_table_separator(table_lines[1]):
        return None

    rows: list[list[str]] = []
    for idx, line in enumerate(table_lines):
        if idx == 1:
            continue
        cells = [clean_cell(cell) for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            value = row[c_idx] if c_idx < len(row) else ""
            table.cell(r_idx, c_idx).text = value


def add_paragraph(doc: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        return
    if stripped.startswith("- "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return
    if stripped[:2].isdigit() and stripped[1:3] == ". ":
        doc.add_paragraph(stripped[3:].strip(), style="List Number")
        return
    doc.add_paragraph(stripped)


def apply_base_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)


def render_markdown(doc: Document, md_path: Path, add_title_as_doc_title: bool = True) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    title_added = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        table_data = parse_table(lines, i)
        if table_data:
            rows, next_index = table_data
            add_table(doc, rows)
            i = next_index
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1 and add_title_as_doc_title and not title_added:
                doc.add_paragraph(text, style="Title")
                title_added = True
            else:
                style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
                doc.add_paragraph(text, style=style)
            i += 1
            continue

        add_paragraph(doc, line)
        i += 1


def export_single(md_name: str) -> Path:
    md_path = SOURCE_DIR / md_name
    doc = Document()
    apply_base_styles(doc)
    render_markdown(doc, md_path)
    output_path = OUTPUT_DIR / f"{md_path.stem}.docx"
    doc.save(output_path)
    return output_path


def export_combined(md_names: Iterable[str]) -> Path:
    doc = Document()
    apply_base_styles(doc)
    first = True
    for md_name in md_names:
        md_path = SOURCE_DIR / md_name
        if not first:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        render_markdown(doc, md_path, add_title_as_doc_title=True)
        first = False
    output_path = OUTPUT_DIR / "servicemen-app-strategy-pack.docx"
    doc.save(output_path)
    return output_path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for md_name in SOURCE_FILES:
        export_single(md_name)
    export_combined(SOURCE_FILES)
    print(f"Exported DOCX files to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
