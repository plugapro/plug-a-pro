#!/usr/bin/env python3
"""Generate a .docx report from the provider funnel JSON data."""

import json
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_kv_table(doc: Document, rows: list[tuple[str, str]], header_color="1F3864"):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        run0 = row.cells[0].paragraphs[0].runs[0]
        run0.bold = True
        run0.font.size = Pt(10)
        run1 = row.cells[1].paragraphs[0].runs[0]
        run1.font.size = Pt(10)
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.0)


def pct(n, d):
    if d == 0:
        return "n/a"
    return f"{n/d*100:.1f}%"


def flag_line(condition: bool, label: str) -> str:
    return f"{'⚠ CRITICAL' if condition else '✓'} {label}"


def generate(json_path: str, out_path: str):
    with open(json_path, "r") as f:
        data = json.load(f)

    report_date = data.get("reportDate", datetime.today().strftime("%Y-%m-%d"))
    apps = data["applications"]
    providers = data["providers"]
    kyc = data["kyc"]
    stuck = data["postApprovalStuck"]
    costs = data["costs"]
    fully_available = data["fullyAvailable"]

    doc = Document()

    # DB error banner
    db_error = data.get("error")
    if db_error:
        p = doc.add_paragraph()
        run = p.add_run(f"⚠ DATA UNAVAILABLE: {db_error}")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # Title
    title = doc.add_heading("Plug A Pro", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Provider Acquisition & Onboarding Funnel Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].bold = True

    date_p = doc.add_paragraph(f"Report date: {report_date}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ── Section 1: Applications ──────────────────────────────────────────────
    add_heading(doc, "1. Application Funnel", level=1)
    approval_rate = apps.get("approvalRate", 0)
    add_kv_table(doc, [
        ("Total applications (all time)", str(apps["total"])),
        ("Submitted in last 30 days", str(apps["last30Days"])),
        ("Pending review", str(apps["pending"])),
        ("Approved", f"{apps['approved']} ({approval_rate:.1f}%)"),
        ("Rejected", str(apps["rejected"])),
        ("Cancelled", str(apps["cancelled"])),
        ("More info required", str(apps["moreInfoRequired"])),
    ])
    doc.add_paragraph()

    # ── Section 2: Provider Onboarding ──────────────────────────────────────
    add_heading(doc, "2. Provider Onboarding Pipeline", level=1)
    live_pct = pct(fully_available, providers["total"])
    add_kv_table(doc, [
        ("Total provider records", str(providers["total"])),
        ("Account created (Supabase auth)", str(providers["withAccount"])),
        ("KYC verified", f"{kyc['verified']} of {kyc['total']}"),
        ("Verified for leads", str(providers["verified"])),
        ("Fully live for leads", f"{fully_available} ({live_pct} of approved)"),
    ])
    doc.add_paragraph()

    # ── Section 3: KYC Breakdown ─────────────────────────────────────────────
    add_heading(doc, "3. KYC Status Breakdown", level=1)
    add_kv_table(doc, [
        ("Not started", str(kyc["notStarted"])),
        ("In progress", str(kyc["inProgress"])),
        ("Submitted (awaiting review)", str(kyc["submitted"])),
        ("Verified", str(kyc["verified"])),
        ("Rejected", str(kyc["rejected"])),
    ])
    doc.add_paragraph()

    # ── Section 4: Post-Approval Stuck ──────────────────────────────────────
    add_heading(doc, "4. Post-Approval Stuck Providers", level=1)
    biggest = stuck.get("biggestDropOff", {})
    add_kv_table(doc, [
        ("Total stuck (approved but not live)", str(stuck["total"])),
        ("No Supabase account created", str(stuck["noAccount"])),
        ("Account created, KYC not started", str(stuck["accountButKycNotStarted"])),
        ("KYC in progress or submitted", str(stuck["kycInProgressOrSubmitted"])),
        ("KYC done, not yet verified", str(stuck["kycDoneNotVerified"])),
        ("Biggest drop-off stage", f"{biggest.get('stage', 'n/a')} ({biggest.get('count', 0)} providers)"),
    ])
    doc.add_paragraph()

    # ── Section 5: Cost Analysis ─────────────────────────────────────────────
    add_heading(doc, "5. Infrastructure Cost per Provider", level=1)
    infra_per = costs.get("infraPerProvider")
    add_kv_table(doc, [
        ("Monthly infrastructure cost (ZAR)", f"R{costs['monthlyInfraZAR']:,.0f}"),
        ("Live providers (denominator)", str(fully_available)),
        ("Infra cost per live provider / month", f"R{infra_per:.2f}" if infra_per else "n/a"),
    ])
    doc.add_paragraph()

    # ── Section 6: Flags ─────────────────────────────────────────────────────
    add_heading(doc, "6. Operational Flags", level=1)
    flags = []
    if fully_available < 20:
        flags.append("CRITICAL: Live provider count below 20 — marketplace capacity at risk")
    if stuck["total"] > 40:
        flags.append("CRITICAL: Post-approval stuck count exceeds 40 — funnel blockage")
    if apps["total"] > 0 and (apps["approved"] / apps["total"]) < 0.70:
        flags.append(
            f"WARNING: Approval rate {approval_rate:.1f}% is below 70% threshold"
        )
    if not flags:
        flags.append("No critical or warning thresholds breached.")

    for flag in flags:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(flag)
        if flag.startswith("CRITICAL"):
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            run.bold = True
        elif flag.startswith("WARNING"):
            run.font.color.rgb = RGBColor(0xD3, 0x85, 0x00)
            run.bold = True
        else:
            run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    doc.add_paragraph()

    # Footer
    footer_p = doc.add_paragraph(
        f"Generated automatically by Claude Code on {report_date}. "
        "Internal use only — do not distribute externally."
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(out_path)
    print(f"Report saved: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: generate_report_docx.py <input.json> <output.docx>")
        sys.exit(1)
    generate(sys.argv[1], sys.argv[2])
